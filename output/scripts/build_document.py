"""Build the revised SID303 Data Exploration DOCX.

Run:  python scripts/build_document.py
Output: ../Tugas_Data_Exploration_SID303_187241037_Revised.docx
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE = Path(__file__).resolve().parent.parent
FIGURES = BASE / "figures"
OUT_DOCX = BASE / "Tugas_Data_Exploration_SID303_187241037_Revised.docx"

NAVY = RGBColor(0x00, 0x00, 0x00)
BLUE = RGBColor(0x00, 0x00, 0x00)
DARK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0x00, 0x00, 0x00)
FONT = "Times New Roman"

TABLE_1A_HEADERS = ["Male", "Age", "Eye Color", "Shoe Size", "Height (in)",
                    "Weight (lb)", "Siblings", "Units", "Handedness"]
TABLE_1A_ROWS = [
    ["1", "20", "Brown", "9.5", "71", "170", "1", "16", "Right"],
    ["0", "19", "Blue", "8", "66", "135", "1", "13", "Right"],
    ["0", "42", "Brown", "7.5", "63", "130", "3", "5", "Right"],
    ["0", "19", "Brown", "8.5", "65", "150", "0", "15", "Left"],
    ["1", "21", "Brown", "11", "70", "185", "5", "19.5", "Right"],
    ["0", "20", "Hazel", "5.5", "60", "105", "2", "11.5", "Right"],
    ["1", "21", "Blue", "12", "76", "210", "2", "9.5", "Right"],
    ["0", "21", "Brown", "10", "70", "140", "0", "8", "Left"],
    ["0", "32", "Brown", "8", "64", "165", "1", "13.5", "Right"],
    ["1", "23", "Brown", "7.5", "63", "145", "6", "12", "Right"],
    ["0", "21", "Brown", "6.5", "61.5", "110", "4", "14", "Right"],
]


def set_run_font(run, size=12, bold=False, color=DARK, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT)


def set_document_page_layout(doc):
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(12)
    normal.font.color.rgb = DARK
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), FONT)

    title = doc.styles["Title"]
    title.font.name = FONT
    title.font.size = Pt(24)
    title.font.bold = True
    title.font.color.rgb = NAVY
    title.paragraph_format.space_after = Pt(6)

    for name, size, before, after in (
        ("Heading 1", 16, 6, 10),
        ("Heading 2", 13, 14, 8),
        ("Heading 3", 12, 10, 6),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = NAVY
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        srpr = style.element.get_or_add_rPr()
        srfonts = srpr.get_or_add_rFonts()
        srfonts.set(qn("w:eastAsia"), FONT)

    caption = doc.styles["Caption"]
    caption.font.name = FONT
    caption.font.size = Pt(9.5)
    caption.font.color.rgb = GRAY
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True


def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    fld1.set(qn("w:dirty"), "true")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-2" \\h \\z \\u '
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Daftar isi akan diperbarui saat dokumen dibuka."
    fld3 = OxmlElement("w:fldChar")
    fld3.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    run._r.append(t)
    run._r.append(fld3)


def shd(fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:color"), "auto")
    el.set(qn("w:fill"), fill)
    return el


def set_table_borders(table, color="D1D5DB"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)
    cell_mar = OxmlElement("w:tblCellMar")
    for side, width in (("top", 40), ("left", 80), ("bottom", 40), ("right", 80)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(width))
        el.set(qn("w:type"), "dxa")
        cell_mar.append(el)
    tblPr.append(cell_mar)


def repeat_header_row(table):
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    trPr.append(tbl_header)


def add_academic_table(doc, headers, rows, widths=None, font_size=9):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        cell._tc.get_or_add_tcPr().append(shd("D9D9D9"))
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)
        run = p.add_run(header)
        set_run_font(run, size=font_size, bold=True, color=DARK)

    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i + 1, j)
            if i % 2 == 1:
                cell._tc.get_or_add_tcPr().append(shd("F2F6FB"))
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)

    if widths:
        table.autofit = False
        tbl = table._tbl
        tblPr = tbl.tblPr
        layout = OxmlElement("w:tblLayout")
        layout.set(qn("w:type"), "fixed")
        tblPr.append(layout)
        grid = tbl.find(qn("w:tblGrid"))
        if grid is not None:
            tbl.remove(grid)
        grid = OxmlElement("w:tblGrid")
        for width in widths:
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(int(width * 567)))
            grid.append(gc)
        tbl.insert(list(tbl).index(tblPr) + 1, grid)
        for i in range(len(rows) + 1):
            for j, width in enumerate(widths):
                table.cell(i, j).width = Cm(width)

    repeat_header_row(table)
    return table


def add_figure(doc, image_path, caption, interpretation=None, width_cm=14.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))

    cap = doc.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_with_next = interpretation is not None

    if interpretation:
        ip = doc.add_paragraph()
        ip.paragraph_format.space_after = Pt(10)
        r = ip.add_run("Interpretasi: ")
        set_run_font(r, size=11, bold=True, color=NAVY)
        r2 = ip.add_run(interpretation)
        set_run_font(r2, size=11)


def add_insight(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    cell._tc.get_or_add_tcPr().append(shd("D9EAF7"))
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:color"), "1F4E78")
    borders.append(left)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("Insight: ")
    set_run_font(r, size=10.5, bold=True, color=NAVY)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5)


def add_question(doc, number, title, question, answer_segments):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(f"{number} {title}")
    set_run_font(r, size=11.5, bold=True, color=NAVY)

    q = doc.add_paragraph()
    q.paragraph_format.keep_with_next = True
    rq = q.add_run(question)
    set_run_font(rq, size=11)

    a = doc.add_paragraph()
    a.paragraph_format.keep_together = True
    for text, bold in answer_segments:
        r = a.add_run(text)
        set_run_font(r, size=11, bold=bold)


def add_section(doc, label, subtitle, source=None):
    h = doc.add_heading(label, level=1)
    h.paragraph_format.page_break_before = True
    sp = doc.add_paragraph()
    sp.paragraph_format.keep_with_next = True
    sp.paragraph_format.space_after = Pt(4)
    r = sp.add_run(subtitle)
    set_run_font(r, size=13, bold=True, color=NAVY)
    if source:
        src = doc.add_paragraph()
        src.paragraph_format.space_after = Pt(10)
        r = src.add_run(source)
        set_run_font(r, size=9.5, italic=True, color=GRAY)


def add_cover(doc):
    # Cover 1 halaman penuh - spacing dihitung untuk mengisi halaman A4 (margin 2.3cm atas/bawah)
    # Konten area: ~25.1cm tinggi. Elemen: judul, logo, nama|NIM, dosen, 5 baris institusi.
    
    # Judul mata kuluk (tanpa kode)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(36)  # ~1.27cm dari margin atas
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("Analisis dan Visualisasi Data")
    set_run_font(r, size=24, bold=True, color=DARK)

    # Logo UNAIR - diperbesar
    logo = Path("/home/misbahul45/Downloads/logo-unair.png")
    if logo.exists():
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_after = Pt(36)
        lp.add_run().add_picture(str(logo), width=Cm(5.5))

    # Nama | NIM
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    r = p.add_run("MISBAHUL MUTTAQIN | 187241037")
    set_run_font(r, size=16, bold=True, color=DARK)

    # Dosen - TERVERIFIKASI dari Cyber Campus
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(48)
    r = p.add_run("Dosen: Dr. Rimuljo Hendradi, S.Si., M.Si.")
    set_run_font(r, size=14, color=DARK)

    # Institusi (5 baris) - size sedikit lebih besar, spacing merata
    inst_lines = [
        "PROGRAM STUDI SISTEM INFORMASI",
        "FAKULTAS SAINS DAN TEKNOLOGI",
        "UNIVERSITAS AIRLANGGA",
        "SURABAYA",
        "2026",
    ]
    for i, line in enumerate(inst_lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(10 if i < 4 else 0)
        r = p.add_run(line)
        set_run_font(r, size=14, bold=(line == "2026"), color=DARK)

    doc.add_page_break()


def add_toc_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("DAFTAR ISI")
    set_run_font(r, size=16, bold=True, color=NAVY)
    add_toc(doc)
    doc.add_page_break()


def add_part1(doc):
    add_section(
        doc,
        "BAGIAN 1",
        "RINGKASAN CHAPTER 1: Introduction to Data Visualization and Visual Data Mining",
        source=("Sumber: Soukup, T., & Davidson, I. (2002). Visual Data Mining: "
                "Techniques and Tools for Data Visualization and Mining. John Wiley & Sons."),
    )

    doc.add_heading("1.1 Overview: Mengapa Visualisasi Data?", level=2)
    doc.add_paragraph(
        "Setiap hari kita menjumpai visualisasi data, misalnya grafik batang untuk "
        "menyampaikan hasil survei kependudukan, grafik garis untuk tren pasar keuangan, "
        "dan peta untuk pola cuaca geografis. Hal ini terjadi karena visualisasi dua dan "
        "tiga dimensi merupakan cara yang paling efektif untuk mengomunikasikan data "
        "dalam jumlah besar yang rumit."
    )
    doc.add_paragraph(
        "Bab ini memperkenalkan dua kelas besar visualisasi. Kelas pertama adalah data "
        "visualization tools and techniques, yaitu alat yang membantu pengguna membuat "
        "gambar dua atau tiga dimensi dari data bisnis sehingga mudah diinterpretasi. "
        "Dalam pendekatan ini manusia bertindak sebagai mesin pengenalan pola (pattern "
        "recognition engine): dengan memeriksa dan berinteraksi dengan visualisasi, "
        "pengguna dapat mengidentifikasi informasi atau pola yang menarik, baik yang "
        "non-trivial, implisit, sebelumnya tidak diketahui, maupun yang berpotensi "
        "berguna. Kelas kedua adalah visual data mining tools and techniques, yaitu alat "
        "yang membantu pengguna membuat visualisasi dari model data mining untuk "
        "memahami pola yang ditemukan algoritma. Pengguna dapat memeriksa dan "
        "berinteraksi dengan visualisasi model prediktif maupun deskriptif untuk "
        "memahami dan memvalidasi pola yang ditemukan, serta mengevaluasi hasil model."
    )
    doc.add_paragraph(
        "Model data mining dapat dipahami sebagai kumpulan generalisasi atau pola dari "
        "data bisnis, yaitu abstraksi dari suatu tugas. Beberapa tool menjelaskan alasan "
        "di balik keputusan, sementara yang lain bersifat black box. Dalam kedua kasus, "
        "visualisasi adalah kunci untuk menemukan pola baru dan mengomunikasikannya "
        "kepada pengambil keputusan. Kombinasi data visualization dan visual data mining "
        "yang efektif memberikan payoff dan ROI yang substansial bagi bisnis. Pemahaman "
        "dasar tentang tipe-tipe visualisasi diperlukan sebelum memulai metodologi "
        "delapan langkah VDM (Visual Data Mining) yang dibahas pada Bab 2 sampai Bab 9."
    )

    add_figure(
        doc,
        FIGURES / "figure_01_visualization_framework.png",
        "Gambar 1. Hubungan data visualization dan visual data mining dalam proses memperoleh insight.",
        interpretation=(
            "Visualisasi pada Gambar 1 merangkum hubungan konseptual kedua pendekatan. "
            "Data visualization memanfaatkan kemampuan manusia mengenali pola secara "
            "langsung dari representasi visual data bisnis, sedangkan visual data mining "
            "menempuh jalur model: pola ditemukan oleh algoritma, divisualisasikan, lalu "
            "diinterpretasi dan divalidasi oleh pengguna. Kedua jalur berakhir pada "
            "pemahaman yang mendukung keputusan."
        ),
    )

    doc.add_heading("1.2 Visualization Data Sets", level=2)
    doc.add_paragraph(
        "Mayoritas data bisnis tersimpan sebagai satu tabel informasi, yaitu sejumlah "
        "kolom berhingga dan satu atau lebih baris data. Contoh sederhana adalah data set "
        "WEATHER dengan kolom CITY, DATE, TEMPERATURE, HUMIDITY, dan CONDITION. Setiap "
        "baris (record) adalah satu fakta data, dan tingkat detail atau granularitas "
        "fakta (unit eksperimen) berada pada level kota. Visualisasi memetakan kolom dan "
        "baris tersebut menjadi gambar dua atau tiga dimensi."
    )

    doc.add_heading("1.3 Visualization Data Types", level=2)
    doc.add_paragraph(
        "Kolom dalam data set bisnis berisi dua tipe nilai. Tipe pertama adalah discrete "
        "(kategorikal), yaitu nilai-nilai berhingga berupa karakter, integer, atau "
        "rentang terkelompok. Jika nilai memiliki urutan inheren, tipe ini disebut "
        "ordinal, misalnya SMALL, MEDIUM, LARGE atau kelompok umur 0 sampai 21 dan 22 "
        "sampai 35. Tipe kedua adalah continuous (numerik atau date), yaitu nilai yang "
        "dapat mengambil rentang penuh dan berpotensi tak hingga, misalnya tanggal, "
        "bilangan presisi ganda, atau floating-point seperti TEMPERATURE, HUMIDITY, dan "
        "TOTAL_SALES."
    )

    table = add_academic_table(
        doc,
        ["Tipe", "Karakteristik", "Contoh"],
        [
            ["Discrete", "Nilai berbentuk kategori atau kelompok tertentu", "CITY, CONDITION"],
            ["Continuous", "Nilai berada dalam rentang numerik", "TEMPERATURE, HUMIDITY"],
        ],
        widths=[3.0, 7.5, 5.5],
    )
    cap = doc.add_paragraph("Tabel 1. Ringkasan tipe data pada data set bisnis.", style="Caption")

    doc.add_heading("1.4 Visual versus Data Dimensions", level=2)
    doc.add_paragraph(
        "Penting untuk membedakan dua istilah. Visual dimension berkaitan dengan sistem "
        "koordinat spasial, yaitu sumbu x, y, dan z, serta atribut seperti warna, "
        "opasitas, tinggi, atau ukuran objek grafis. Data dimension berkaitan dengan "
        "jumlah kolom dalam data set bisnis. Untuk membuat visualisasi, kolom yang "
        "diselidiki dipilih dari data set menjadi graphical data table, yang memetakan "
        "nilai kolom ke titik data pada sistem koordinat. Sebagai contoh, grafik kolom "
        "membandingkan TEMPERATURE dan HUMIDITY (data dimensions kontinu) berdasarkan "
        "CITY (data dimension diskret), dengan tinggi batang mewakili nilai kolom."
    )

    doc.add_heading("1.5 Data Visualization Tools", level=2)
    doc.add_paragraph(
        "Tool visualisasi data diklasifikasikan menjadi dua kategori utama. Kategori "
        "pertama membandingkan nilai satu kolom dengan kolom lain menggunakan sistem "
        "koordinat spasial, sedangkan kategori kedua mengeksploitasi struktur inheren "
        "dari data."
    )
    doc.add_heading("A. Multidimensional Data Visualization Tools", level=3)
    doc.add_paragraph(
        "Grafik kolom dan bar (column dan bar graphs) membandingkan nilai kontinu "
        "melintasi nilai diskret; stacked column atau bar mengakumulasi nilai sehingga "
        "batang bertumpuk. Melalui inspeksi visual, aturan seperti suhu cenderung lebih "
        "tinggi daripada kelembapan kecuali saat hujan dapat ditemukan. Distribution dan "
        "histogram graphs menampilkan proporsi nilai: distribution digunakan untuk kolom "
        "diskret, sedangkan histogram (frequency graph) untuk kolom kontinu. Keduanya "
        "berguna untuk mendeteksi ketidakseimbangan data dan skewness."
    )
    doc.add_paragraph(
        "Box graphs (box plot) menampilkan statistik deskriptif kolom kontinu, yaitu "
        "kuartil 25 persen dan 75 persen dengan panjang bar sebagai ukuran variabilitas, "
        "serta minimum, maksimum, median, mean, dan standar deviasi; posisi median "
        "mengindikasikan skewness. Line graphs menunjukkan tren deret waktu, dengan "
        "variasi high-low-close untuk tren saham dan radar graph pada koordinat 360 "
        "derajat. Satu visualisasi garis dapat mengomunikasikan ribuan informasi, "
        "misalnya lebih dari 4.500 titik data obligasi, yang sulit dilihat pada laporan "
        "tabular green-bar."
    )
    doc.add_paragraph(
        "Scatter graphs memetakan setiap baris ke titik pada koordinat x-y (atau 3-D) "
        "untuk menyelidiki hubungan antar kolom kontinu; variasi bubble graph "
        "menambahkan dimensi ukuran objek. Pie dan doughnut graphs menampilkan kontribusi "
        "tiap nilai terhadap total, dengan doughnut graph yang dapat membandingkan "
        "beberapa kolom kontinu sekaligus."
    )
    doc.add_heading("B. Hierarchical and Landscape Data Visualization Tools", level=3)
    doc.add_paragraph(
        "Tree visualizations menampilkan data sebagai pohon; setiap level bercabang "
        "berdasarkan nilai atribut yang berbeda, dan node berisi agregasi seperti jumlah, "
        "rata-rata, atau count dalam bentuk batang atau disk. Contohnya adalah proporsi "
        "keluarga penerima Medicaid berdasarkan tipe keluarga dan region. Map "
        "visualizations menampilkan nilai kolom sebagai elemen grafis pada peta "
        "berdasarkan kunci spasial atau geografis, misalnya jumlah registrasi akun baru "
        "per negara bagian yang diwarnai sesuai jumlahnya."
    )

    doc.add_heading("1.6 Visual Data Mining Tools", level=2)
    doc.add_paragraph(
        "Tool visual data mining membuat gambar dua atau tiga dimensi tentang bagaimana "
        "model data mining mengambil keputusan. Decision tree divisualisasikan sebagai "
        "hierarchical tree graph sehingga struktur model mudah dipahami, misalnya untuk "
        "memprediksi potensi gaji. Tidak semua algoritma mudah divisualisasikan; neural "
        "network dengan unit pemrosesan yang saling terhubung dalam lapisan input, "
        "hidden, dan output masih menjadi pertanyaan riset aktif."
    )
    doc.add_paragraph(
        "Gains chart adalah grafik garis yang membandingkan performa model dalam "
        "memprediksi kejadian target terhadap tebakan acak. Cumulative gain menyatakan "
        "proporsi seluruh kejadian target hingga persentil tertentu. Alat ini berguna "
        "untuk membandingkan model, memantau performa setelah deployment, dan mendeteksi "
        "model yang stale."
    )

    doc.add_heading("1.7 Summary", level=2)
    doc.add_paragraph(
        "Chapter 1 merangkum tool dan teknik data visualization serta visual data mining "
        "untuk menemukan tren, perilaku, dan anomali yang sebelumnya tidak diketahui "
        "dalam data bisnis. Alur yang dibangun dimulai dari data, dilanjutkan ke "
        "representasi visual, interpretasi manusia, pengenalan pola, hingga insight yang "
        "mendukung keputusan. Pada jalur visual data mining, model data mining "
        "divisualisasikan agar dapat diinterpretasi dan divalidasi. Bab 2 sampai Bab 9 "
        "menyajikan metodologi delapan langkah VDM yang terbukti untuk membangun solusi "
        "business intelligence."
    )


def add_part2(doc):
    add_section(
        doc,
        "BAGIAN 2",
        "LATIHAN SECTION 1.2: Soal 1.1 sampai 1.10 (hal. 51)",
        source=("Sumber: Gould, R., Ryan, C., & Wong, R. (2017). Essential Statistics: "
                "Exploring the World through Data (2nd ed.). Pearson Education Limited."),
    )

    doc.add_paragraph(
        "Data pada Tabel 1A dikumpulkan dari sebuah kelas olahraga (gym class). Kepala "
        "kolom menunjukkan variabel yang dicatat, dan setiap baris mewakili satu individu "
        "dalam kelas."
    )
    cap = doc.add_paragraph("Tabel 1A. Data kelas olahraga dengan 11 observasi dan 9 variabel.", style="Caption")
    add_academic_table(
        doc,
        TABLE_1A_HEADERS,
        TABLE_1A_ROWS,
        widths=[1.1, 1.1, 1.9, 1.9, 2.1, 2.1, 1.6, 1.3, 2.4],
    )

    doc.add_heading("Ringkasan Eksplorasi Data", level=2)
    doc.add_paragraph(
        "Tabel 1A memuat 11 observasi dengan 9 variabel. Eye Color didominasi Brown "
        "(8 dari 11 observasi), Handedness didominasi Right (9 dari 11 observasi), dan "
        "berdasarkan batas 12 units, 7 mahasiswa tergolong full-time serta 4 mahasiswa "
        "part-time. Scatter plot tinggi terhadap berat badan memperlihatkan kecenderungan "
        "asosiasi positif pada sampel ini. Bagian ini hanya menggambarkan eksplorasi "
        "deskriptif; tidak dilakukan generalisasi ke populasi karena data berasal dari "
        "satu kelas."
    )

    add_figure(
        doc,
        FIGURES / "figure_02_eye_color_distribution.png",
        "Gambar 2. Distribusi warna mata pada 11 observasi dalam Tabel 1A.",
        interpretation=(
            "Brown merupakan kategori warna mata yang paling banyak ditemukan, yaitu 8 "
            "dari 11 observasi. Blue muncul pada 2 observasi, sedangkan Hazel hanya "
            "ditemukan pada 1 observasi."
        ),
    )
    add_insight(doc, "Sebanyak 8 dari 11 observasi memiliki eye color Brown.")

    add_figure(
        doc,
        FIGURES / "figure_04_height_weight_scatter.png",
        "Gambar 4. Hubungan antara tinggi dan berat badan pada data Tabel 1A.",
        interpretation=(
            "Scatter plot menunjukkan kecenderungan hubungan positif antara tinggi dan "
            "berat badan pada sampel ini. Nilai korelasi Pearson yang dihitung dari 11 "
            "observasi adalah r = 0,853, yang menunjukkan hubungan positif yang cukup "
            "kuat. Namun, ukuran sampel relatif kecil sehingga visualisasi ini digunakan "
            "sebagai eksplorasi data, bukan dasar untuk melakukan generalisasi populasi."
        ),
    )

    add_question(
        doc, "1.1", "Variables",
        "In Table 1A, how many variables are there?",
        [("Jawaban: ", True),
         ("Terdapat 9 variabel, yaitu Male, Age, Eye Color, Shoe Size, Height (in), "
          "Weight (lb), Siblings, Units, dan Handedness.", False)],
    )

    add_question(
        doc, "1.2", "People",
        "In Table 1A, there are observations on how many people?",
        [("Jawaban: ", True),
         ("Terdapat 11 orang, sesuai dengan 11 baris observasi pada Tabel 1A.", False)],
    )

    add_question(
        doc, "1.3", "(TRY, Example 1)",
        "Are the following variables, from Table 1A, numerical or categorical? Explain. "
        "a. Handedness  b. Age",
        [("Jawaban: ", True),
         ("a. Handedness termasuk variabel kategorikal karena nilai Right dan Left "
          "menunjukkan kelompok, bukan besaran numerik yang dapat dihitung secara "
          "aritmetika. b. Age termasuk variabel numerik karena umur dinyatakan sebagai "
          "besaran terukur dalam tahun dan dapat dibandingkan secara kuantitatif.", False)],
    )

    add_question(
        doc, "1.4", "Numerical or Categorical",
        "Are the following variables, from Table 1A, numerical or categorical? Explain. "
        "a. Shoe size  b. Eye color",
        [("Jawaban: ", True),
         ("a. Shoe size termasuk variabel numerik karena ukuran sepatu diukur pada skala "
          "numerik (9.5, 8, 7.5, dan seterusnya) dan dapat dibandingkan secara "
          "kuantitatif. b. Eye color termasuk variabel kategorikal karena nilai Brown, "
          "Blue, dan Hazel menunjukkan kategori tanpa urutan atau makna numerik.", False)],
    )

    add_question(
        doc, "1.5", "Another Numerical Variable",
        "Give an example of another numerical variable we might have recorded for the "
        "students whose data are in Table 1A.",
        [("Jawaban: ", True),
         ("Contoh variabel numerik lain yang mungkin dicatat adalah GPA (indeks "
          "prestasi), nilai ujian, jam belajar per minggu, jumlah kelas yang dihadiri, "
          "atau tinggi badan dalam sentimeter.", False)],
    )

    add_question(
        doc, "1.6", "Another Categorical Variable",
        "Give an example of another categorical variable we might have recorded for the "
        "students whose data are in Table 1A.",
        [("Jawaban: ", True),
         ("Contoh variabel kategorikal lain adalah program studi atau jurusan, jenis "
          "kelamin (jika tidak dikodekan), status tempat tinggal (asrama atau luar), "
          "atau olahraga favorit.", False)],
    )

    add_question(
        doc, "1.7", "Coding",
        "Suppose you decided to code eye color using 1 for Brown Eyes and 0 for Not "
        "Brown Eyes. What would be the label at the top of the column, and how many ones "
        "and zeros would there be?",
        [("Jawaban: ", True),
         ("Label kolom adalah Brown, dengan kode 1 untuk Brown Eyes dan 0 untuk Not "
          "Brown Eyes. Angka 1 dan 0 pada variabel kategorikal merupakan kode kategori, "
          "bukan nilai numerik. Dari Tabel 1A terdapat 8 observasi bermata Brown dan 3 "
          "observasi bukan Brown (Blue 2 dan Hazel 1), sehingga kolom berisi 8 angka 1 "
          "dan 3 angka 0.", False)],
    )

    add_question(
        doc, "1.8", "Coding",
        "Suppose you decided to code handedness using Right-handed as the label for the "
        "column. How many ones and how many zeros would there be?",
        [("Jawaban: ", True),
         ("Label kolom adalah Right-handed, dengan kode 1 untuk Right dan 0 untuk Left. "
          "Dari Tabel 1A terdapat 9 observasi Right dan 2 observasi Left, sehingga "
          "kolom berisi 9 angka 1 dan 2 angka 0.", False)],
    )

    add_question(
        doc, "1.9", "Coding",
        "Explain why the variable Male, in Table 1A, is categorical, even though its "
        "values are numbers. Often, it does not make sense, or is not even possible, to "
        "add the values of a categorical variable. Does it make sense for Male? If so, "
        "what does the sum represent?",
        [("Jawaban: ", True),
         ("Variabel Male bersifat kategorikal karena angka 1 dan 0 hanyalah kode untuk "
          "kategori (1 untuk laki-laki dan 0 untuk bukan laki-laki), bukan besaran yang "
          "dapat diukur. Menjumlahkan nilai Male masuk akal dalam kasus ini karena setiap "
          "kode 1 mewakili satu orang laki-laki: 1 + 0 + 0 + 0 + 1 + 0 + 1 + 0 + 0 + 1 + "
          "0 = 4. Jumlah tersebut merepresentasikan banyaknya laki-laki di kelas, yaitu "
          "4 dari 11 orang.", False)],
    )

    add_question(
        doc, "1.10", "Coding",
        "Students with fewer than 12 units in the current term are considered part-time. "
        "Create a new categorical variable that classifies each student in Table 1A as "
        "full-time (12 or more units) or part-time. Call this variable Full. Report the "
        "values in a column in the same order as those in the table. Use codes (1 and 0) "
        "in your column.",
        [("Jawaban: ", True),
         ("Variabel Full menggunakan kode 1 untuk full-time (12 units atau lebih) dan 0 "
          "untuk part-time (kurang dari 12 units). Kolom nilai sesuai urutan Tabel 1A "
          "adalah 1, 1, 0, 1, 1, 0, 0, 0, 1, 1, 1. Dengan demikian terdapat 7 mahasiswa "
          "full-time (units 16, 13, 15, 19.5, 13.5, 12, dan 14) serta 4 mahasiswa "
          "part-time (units 5, 11.5, 9.5, dan 8).", False)],
    )

    add_figure(
        doc,
        FIGURES / "figure_03_fulltime_distribution.png",
        "Gambar 3. Distribusi status full-time dan part-time berdasarkan jumlah units.",
        interpretation=(
            "Berdasarkan batas 12 units, 7 dari 11 mahasiswa dikategorikan sebagai "
            "full-time dan 4 mahasiswa sebagai part-time. Visualisasi ini merupakan "
            "representasi langsung dari jawaban soal 1.10."
        ),
    )


def add_part3(doc):
    add_section(
        doc,
        "BAGIAN 3",
        "LATIHAN SECTION 1.3: Soal 1.15 (hal. 52)",
        source=("Sumber: Gould, R., Ryan, C., & Wong, R. (2017). Essential Statistics: "
                "Exploring the World through Data (2nd ed.). Pearson Education Limited."),
    )

    doc.add_paragraph(
        "TRY 1.15 Older Siblings (Example 3): At a small four-year college, some "
        "psychology students were asked whether or not they had at least one older "
        "sibling. The table shows the results for men and women and shows some of the "
        "totals."
    )

    cap = doc.add_paragraph("Tabel 2. Hasil survei older sibling sebelum total dilengkapi.", style="Caption")
    add_academic_table(
        doc,
        ["", "Men", "Women", "Total"],
        [
            ["Yes, Older Sibling", "12", "55", "?"],
            ["No Older Sibling", "11", "39", "50"],
            ["Total", "23", "?", "117"],
        ],
        widths=[4.6, 3.2, 3.2, 3.2],
    )

    add_question(
        doc, "a.", "Totals yang Belum Ditampilkan",
        "Calculate the totals that are not shown, and report them in the table.",
        [("Jawaban: ", True),
         ("Total Yes, Older Sibling = 12 + 55 = 67, dan Total Women = 55 + 39 = 94. "
          "Verifikasi: 23 + 94 = 117 dan 67 + 50 = 117, sehingga kedua total konsisten "
          "dengan total keseluruhan 117.", False)],
    )

    cap = doc.add_paragraph("Tabel 3. Hasil survei older sibling setelah total dilengkapi.", style="Caption")
    add_academic_table(
        doc,
        ["", "Men", "Women", "Total"],
        [
            ["Yes, Older Sibling", "12", "55", "67"],
            ["No Older Sibling", "11", "39", "50"],
            ["Total", "23", "94", "117"],
        ],
        widths=[4.6, 3.2, 3.2, 3.2],
    )

    add_question(
        doc, "b.", "Persentase Men dengan Older Sibling",
        "What percentage of the men had an older sibling?",
        [("Jawaban: ", True),
         ("12/23 × 100% = 52,17%.", False)],
    )

    add_question(
        doc, "c.", "Persentase Men tanpa Older Sibling",
        "What percentage of the men did not have an older sibling?",
        [("Jawaban: ", True),
         ("11/23 × 100% = 47,83%.", False)],
    )

    add_question(
        doc, "d.", "Persentase Women dengan Older Sibling",
        "What percentage of the women had an older sibling?",
        [("Jawaban: ", True),
         ("55/94 × 100% = 58,51%.", False)],
    )

    add_question(
        doc, "e.", "Persentase Seluruh Responden dengan Older Sibling",
        "What percentage of the people had an older sibling?",
        [("Jawaban: ", True),
         ("67/117 × 100% = 57,27%.", False)],
    )

    add_question(
        doc, "f.", "Proporsi Women di antara Responden dengan Older Sibling",
        "What percentage of the people with an older sibling were women?",
        [("Jawaban: ", True),
         ("55/67 × 100% = 82,09%.", False)],
    )

    add_question(
        doc, "g.", "Estimasi pada 600 Women",
        "Suppose that in a group of 600 women, the percentage who have an older sibling "
        "is the same as in the sample here. How many of the 600 women would have an "
        "older sibling?",
        [("Jawaban: ", True),
         ("600 × (55/94) = 351,06, sehingga diperkirakan sekitar 351 wanita yang memiliki "
          "older sibling (dibulatkan ke bilangan bulat terdekat).", False)],
    )

    add_figure(
        doc,
        FIGURES / "figure_05_older_sibling_gender.png",
        "Gambar 5. Perbandingan mahasiswa yang memiliki dan tidak memiliki older sibling berdasarkan gender.",
        interpretation=(
            "Secara jumlah absolut, women mendominasi kedua kategori karena jumlah women "
            "dalam sampel lebih besar. Jika dibandingkan berdasarkan proporsi di dalam "
            "masing-masing gender, 58,51% women memiliki older sibling, sedangkan pada "
            "men proporsinya sebesar 52,17%. Perbedaan ini menunjukkan pentingnya "
            "membedakan perbandingan jumlah absolut dengan perbandingan persentase dalam "
            "kelompok."
        ),
    )


def add_references(doc):
    h = doc.add_heading("DAFTAR PUSTAKA", level=1)
    h.paragraph_format.page_break_before = True

    refs = [
        "Soukup, T., & Davidson, I. (2002). Visual Data Mining: Techniques and Tools for "
        "Data Visualization and Mining. John Wiley & Sons.",
        "Gould, R., Ryan, C., & Wong, R. (2017). Essential Statistics: Exploring the "
        "World through Data (2nd ed.). Pearson Education Limited.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.75)
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(ref)
        set_run_font(r, size=11)


def generate_document():
    doc = Document()
    set_document_page_layout(doc)
    configure_styles(doc)

    add_cover(doc)
    add_toc_page(doc)
    add_part1(doc)
    add_part2(doc)
    add_part3(doc)
    add_references(doc)

    doc.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX}")


if __name__ == "__main__":
    generate_document()